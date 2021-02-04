import copy
import re
import string
from collections import defaultdict

import docx
import pandas as pd
import spacy


class DialogueParser:
    @staticmethod
    def split_speaker_text(text_line):
        """
        Splits Speaker and Text of paragraph
        :type text_line: str
        :param text_line: Paragraph to split
        :return: tuple of (speaker, text), if no speaker: (None, Text)
        """
        pattern = re.compile(r"^(([AB])[:;] )(.*)")
        matches = pattern.match(text_line)

        if matches:
            return matches.group(2), matches.group(3)
        else:
            return None, None

    @staticmethod
    def clean_comments(text: str) -> str:
        """
        Remove Comments from Text
        :type text: str
        :param text: text to remove comments from
        :return: cleaned text
        """
        if text:
            text = re.sub(r"\(.*?\)", "", text)
            text = re.sub(r"\[.*?\]", "", text)
            # Remove Digits and Symbols
            # Remove whitespace
            text = re.sub(' +', ' ', text)
            text = re.sub(r' \.', '.', text)

            return text
        else:
            return " "

    def split_sentences(self, text, using_punctuation=True):
        if not using_punctuation:
            text = ''.join([symb for symb in text if not (
                symb.isdigit() or symb in string.punctuation)])
        doc = self.nlp(text)

        sentences = [sent.text for sent in doc.sents]

        return sentences

    def __init__(self, file, nlp, clean_comments=True, genders=None) -> None:
        """
        Create Parser Object
        :type file: str
        :type nlp: callable
        :type clean_comments: Bool
        :param file: Doc-Object to parse
        :param nlp: NLP Object to use to split sentences
        :param clean_comments: clean comments like "Unverständlich", default: True
        """
        self.document = docx.Document(file)

        self.nlp = nlp
        if not self.nlp.has_pipe("sentencizer"):
            self.nlp.add_pipe(self.nlp.create_pipe('sentencizer'))

        dictionary_rows = []
        sentence_dict_row = []
        self.paragraphs = [paragraph.text for paragraph in self.document.paragraphs]
        self.genders = {}
        if genders:
            self.genders["A"], self.genders["B"] = genders
        else:
            self.genders["A"] = self._extract_gender(self.paragraphs[1])
            self.genders["B"] = self._extract_gender(self.paragraphs[2])

        print(self.genders)

        for paragraph_position, paragraph in enumerate(
                list(self.document.paragraphs)[3:]):
            dictionary_row = {}
            dictionary_row["paragraph_position"] = paragraph_position
            speaker, text = self.split_speaker_text(paragraph.text)

            if clean_comments:
                text = DialogueParser.clean_comments(text)

            if speaker:
                dictionary_row["speaker"] = speaker
                dictionary_row["gender"] = self.genders[speaker]
                dictionary_row["raw_text"] = text
                dictionary_row["is_paragraph"] = True
                sentences = self.split_sentences(text, True)
                for index, sentence in enumerate(sentences):
                    sentence_row = copy.copy(dictionary_row)
                    sentence_row["raw_text"] = sentence
                    sentence_row["sentence_index"] = index
                    sentence_dict_row.append(sentence_row)
            else:
                dictionary_row["raw_text"] = paragraph.text
                dictionary_row["is_paragraph"] = False
            dictionary_rows.append(dictionary_row)

        self.paragraph_frame = pd.DataFrame(dictionary_rows)
        self.sentence_frame = pd.DataFrame(sentence_dict_row)
        super().__init__()

        self.df_paragraphs = pd.DataFrame()

    def _extract_gender(self, paragraph):
        without_prefix = re.sub("[AB]: ", "", paragraph)
        without_prefix = re.sub("Er", "M", without_prefix)
        without_prefix = re.sub("Sie", "W", without_prefix)
        return f"{without_prefix}"

    def get_sentences(self, query=None):
        """
        Get sentences of Dialogue as Dataframe, Meta-Data and raw text
        :type query: additional filter for sentences to get, example: get all paragraphs of speaker A
        """
        if query:
            return self.sentence_frame.query(query)
        return self.sentence_frame

    def get_paragraphs(self, query: string = None):
        """
        Get paragraphs of Dialogue as Dataframe, Meta-Data and raw text
        :param query: additional filter for paragraphs to get, example: get all paragraphs of speaker A
        """
        if query:
            return self.paragraph_frame.query(query)
        return self.paragraph_frame

    def get_fulltext(self, split_speakers=True):
        if split_speakers:
            return self.paragraph_frame.groupby(by="speaker").agg({"raw_text": ''.join})
        else:
            return self.paragraph_frame["raw_text"].str.cat(sep=" ")


if __name__ == '__main__':
    dp = DialogueParser(r"../data/Paar 105_T1_IM_FW.docx",
                        spacy.load("de_core_news_md"),
                        clean_comments=True)
    sentences = dp.get_sentences()
    fulltext = dp.get_fulltext(False)
    sentences_a = dp.get_sentences(query="speaker=='A'")
    sentences_b = dp.get_sentences(query="speaker=='B'")
    print(sentences_b)
    pass
